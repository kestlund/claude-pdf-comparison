#!/usr/bin/env python3
"""
Document Comparison Tool for Claude API
Compares a group of PDF and Word files against criteria in a reference PDF.
"""

import os
import json
import csv
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
import PyPDF2
import fitz  # PyMuPDF for better text extraction
import docx  # python-docx for Word documents
import mammoth  # Alternative Word processor for better formatting
from anthropic import Anthropic
import argparse
import logging
import pandas as pd
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class ComparisonResult:
    """Structure for storing comparison results"""
    filename: str
    file_path: str
    file_id: str
    meets_criteria: bool
    likelihood: float
    key_findings: List[str]

class DocumentProcessor:
    """Handles PDF and Word document text extraction and preprocessing"""
    
    def __init__(self):
        self.max_chars = 50000  # Limit for API calls
    
    def extract_text_pymupdf(self, pdf_path: str) -> str:
        """Extract text using PyMuPDF (better for complex layouts)"""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed for {pdf_path}: {e}")
            return ""
    
    def extract_text_pypdf2(self, pdf_path: str) -> str:
        """Fallback text extraction using PyPDF2"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
            return text.strip()
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed for {pdf_path}: {e}")
            return ""
    
    def extract_text_word_docx(self, doc_path: str) -> str:
        """Extract text from Word document using python-docx"""
        try:
            doc = docx.Document(doc_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"python-docx extraction failed for {doc_path}: {e}")
            return ""
    
    def extract_text_word_mammoth(self, doc_path: str) -> str:
        """Extract text from Word document using mammoth (better formatting)"""
        try:
            with open(doc_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value.strip()
        except Exception as e:
            logger.error(f"Mammoth extraction failed for {doc_path}: {e}")
            return ""
    
    def extract_text_word(self, doc_path: str) -> str:
        """Extract text from Word document with fallback methods"""
        # Try mammoth first (better formatting preservation)
        text = self.extract_text_word_mammoth(doc_path)
        
        # Fallback to python-docx if needed
        if not text or len(text) < 100:
            text = self.extract_text_word_docx(doc_path)
        
        return text
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or Word document"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_word(file_path)
        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return ""
    
    def extract_text_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with fallback methods"""
        # Try PyMuPDF first (better quality)
        text = self.extract_text_pymupdf(pdf_path)
        
        # Fallback to PyPDF2 if needed
        if not text or len(text) < 100:
            text = self.extract_text_pypdf2(pdf_path)
        
        # Truncate if too long for API
        if len(text) > self.max_chars:
            text = text[:self.max_chars] + "\n\n[TEXT TRUNCATED]"
        
        return text
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract document metadata"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_pdf_metadata(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_word_metadata(file_path)
        else:
            return {'title': '', 'pages': 0, 'file_size': os.path.getsize(file_path)}
    
    def extract_pdf_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """Extract PDF metadata"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                metadata = reader.metadata
                return {
                    'title': metadata.get('/Title', '') if metadata else '',
                    'pages': len(reader.pages),
                    'file_size': os.path.getsize(pdf_path)
                }
        except Exception as e:
            logger.error(f"Metadata extraction failed for {pdf_path}: {e}")
            return {'title': '', 'pages': 0, 'file_size': 0}
    
    def extract_word_metadata(self, doc_path: str) -> Dict[str, Any]:
        """Extract Word document metadata"""
        try:
            if doc_path.endswith('.docx'):
                doc = docx.Document(doc_path)
                core_props = doc.core_properties
                return {
                    'title': core_props.title or '',
                    'pages': 0,  # Word page count is complex to determine
                    'file_size': os.path.getsize(doc_path)
                }
            else:
                # For .doc files, we can't easily extract metadata
                return {
                    'title': '',
                    'pages': 0,
                    'file_size': os.path.getsize(doc_path)
                }
        except Exception as e:
            logger.error(f"Word metadata extraction failed for {doc_path}: {e}")
            return {'title': '', 'pages': 0, 'file_size': os.path.getsize(doc_path)}

class ClaudeComparator:
    """Handles comparisons using Claude API"""
    
    def __init__(self, api_key: str):
        self.client = Anthropic(api_key=api_key)
        self.model = "claude-3-sonnet-20240229"
    
    def create_comparison_prompt(self, criteria_text: str, document_text: str, filename: str) -> str:
        """Create the prompt for document comparison using only criteria from PDF"""
        
        return f"""You are a research compliance expert at a large university.

You are tasked with evaluating research project files for gain of function research criteria.

HERE ARE YOUR INSTRUCTIONS:
1. First, review the criteria for gain of function research as described in the following content:

REFERENCE CRITERIA DOCUMENT:
{criteria_text}

Now, you will analyze the content of each document.

DOCUMENT TO ANALYZE:
Filename: {filename}
Content: {document_text}

To evaluate this document file against the criteria, follow these steps:

1. Carefully read through the entire document content.
2. Compare the content to each point in the criteria for gain of function research.
3. Look for specific mentions of techniques, methodologies, or results that align with the criteria.
4. Pay attention to any disclaimers or statements about the nature of the research.

To determine if the document matches the criteria and calculate the likelihood:

1. Assess whether the research described in the document meets the definition of gain of function research as per the criteria.
2. Evaluate how many of the specific criteria points are met by the research described.
3. Consider the clarity and directness of the language used in describing the research.
4. Determine a likelihood score on a scale of 0 to 100, where 0 means definitely not gain of function research, and 100 means definitely gain of function research.

Identify and summarize key findings that support your evaluation. These should be brief but specific points from the document that relate to the criteria.

After your analysis, provide your results in the following JSON format:

{{
    "meets_criteria": true/false,
    "likelihood": 0-100,
    "key_findings": ["List of key findings"]
}}

Ensure that your evaluation is objective and based solely on the content of the document and the provided criteria. Do not make assumptions beyond what is explicitly stated in the document."""

    async def compare_document(self, criteria_text: str, document_text: str, filename: str,
                             file_path: str, metadata: Dict[str, Any]) -> ComparisonResult:
        """Compare a single document against criteria extracted from criteria PDF"""
        try:
            prompt = self.create_comparison_prompt(criteria_text, document_text, filename)
            
            response = self.client.messages.create(
                model=self.model,
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}]
            )
            
            # Parse the JSON response
            response_text = response.content[0].text
            
            # Extract JSON from response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            json_str = response_text[json_start:json_end]
            
            result_data = json.loads(json_str)
                
            return ComparisonResult(
                filename=filename,
                file_path=file_path,
                file_id="",  # Will be set later
                meets_criteria=result_data.get('meets_criteria', False),
                likelihood=float(result_data.get('likelihood', 0)),
                key_findings=result_data.get('key_findings', [])
            )
            
        except Exception as e:
            logger.error(f"Comparison failed for {filename}: {e}")
            
            return ComparisonResult(
                filename=filename,
                file_path=file_path,
                file_id="",  # Will be set later
                meets_criteria=False,
                likelihood=0.0,
                key_findings=[]
            )

class DocumentComparisonTool:
    """Main tool for comparing documents"""
    
    def __init__(self, api_key: str):
        self.document_processor = DocumentProcessor()
        self.comparator = ClaudeComparator(api_key)
    
    def extract_file_id(self, filename: str) -> str:
        """Extract the first 7 characters of the filename as ID"""
        return filename[:7]
    
    def load_criteria(self, criteria_pdf_path: str) -> str:
        """Load criteria from PDF"""
        logger.info(f"Loading criteria from {criteria_pdf_path}")
        criteria_text = self.document_processor.extract_text(criteria_pdf_path)
        
        if not criteria_text:
            raise ValueError(f"Could not extract text from criteria PDF: {criteria_pdf_path}")
        
        return criteria_text
    
    def find_document_files(self, directory: str) -> List[str]:
        """Find all PDF and Word files in directory"""
        document_files = []
        
        # Find PDF files
        for file_path in Path(directory).glob("*.pdf"):
            document_files.append(str(file_path))
        
        # Find Word files (.docx and .doc)
        for file_path in Path(directory).glob("*.docx"):
            document_files.append(str(file_path))
        for file_path in Path(directory).glob("*.doc"):
            document_files.append(str(file_path))
        
        return sorted(document_files)
    
    async def compare_documents(self, criteria_pdf_path: str, document_directory: str,
                          csv_output_path: str = None, json_output_path: str = None) -> List[ComparisonResult]:
        """Compare all documents in directory against criteria from criteria PDF"""
        
        # Load criteria
        criteria_text = self.load_criteria(criteria_pdf_path)
        
        # Find document files to compare
        document_files = self.find_document_files(document_directory)
        
        if not document_files:
            raise ValueError(f"No PDF or Word files found in directory: {document_directory}")
        
        logger.info(f"Found {len(document_files)} document files to compare")
        logger.info(f"Using criteria from: {criteria_pdf_path}")
        
        results = []
        
        for doc_path in document_files:
            logger.info(f"Processing {doc_path}")
            
            # Extract text and metadata from document
            document_text = self.document_processor.extract_text(doc_path)
            metadata = self.document_processor.extract_metadata(doc_path)
            
            if not document_text:
                logger.warning(f"Could not extract text from {doc_path}")
                continue
            
            # Compare against criteria
            filename = Path(doc_path).name
            file_id = self.extract_file_id(filename)
            result = await self.comparator.compare_document(
                criteria_text, document_text, filename, doc_path, metadata
            )
            
            # Update result with file_id
            result.file_id = file_id
            
            results.append(result)
            logger.info(f"Completed analysis for {filename} - Likelihood: {result.likelihood:.2f}")
        
        # Save results
        if csv_output_path and results:
            self.save_results_csv(results, csv_output_path)
        
        if json_output_path:
            self.save_results_json(results, json_output_path)
        
        return results
    
    def save_results_csv(self, results: List[ComparisonResult], output_path: str):
        """Save results to CSV file"""
        
        # Prepare CSV data
        csv_data = []
        
        for result in results:
            row = {
                'id': result.file_id,
                'filename': result.filename,
                'file_path': result.file_path,
                'meets_criteria': result.meets_criteria,
                'likelihood': round(result.likelihood, 3),
                'key_findings': '; '.join(result.key_findings)
            }
            csv_data.append(row)
        
        # Write CSV file
        if csv_data:
            df = pd.DataFrame(csv_data)
            df.to_csv(output_path, index=False)
            logger.info(f"CSV results saved to {output_path}")
    
    def save_results_json(self, results: List[ComparisonResult], output_path: str):
        """Save results to JSON file"""
        json_data = [asdict(result) for result in results]
        with open(output_path, 'w') as f:
            json.dump(json_data, f, indent=2)
        logger.info(f"JSON results saved to {output_path}")
    
    def print_summary(self, results: List[ComparisonResult]):
        """Print a summary of the results"""
        if not results:
            print("No results to display.")
            return
        
        print(f"\n📊 COMPARISON SUMMARY")
        print(f"{'='*50}")
        print(f"Total files analyzed: {len(results)}")
        
        meets_criteria = sum(1 for r in results if r.meets_criteria)
        print(f"Files meeting criteria: {meets_criteria}")
        print(f"Files not meeting criteria: {len(results) - meets_criteria}")
        
        if results:
            avg_likelihood = sum(r.likelihood for r in results) / len(results)
            print(f"Average likelihood: {avg_likelihood:.2f}")
        
        print(f"\n📋 DETAILED RESULTS:")
        for result in results:
            status = "✅ MEETS" if result.meets_criteria else "❌ DOES NOT MEET"
            print(f"{result.filename}: {status} (Likelihood: {result.likelihood:.1f}%)")

def main():
    """Main function for command-line usage"""
    
    # ADD THESE LINES TO HARDCODE PATHS:
    CRITERIA_PDF_PATH = "./criteria.pdf"
    PDF_DIRECTORY_PATH = "./usda-pdfs"
    
    parser = argparse.ArgumentParser(description="Compare PDF files against criteria with CSV output")
    # Modify these lines to use default paths:
    parser.add_argument("--criteria-pdf", default=CRITERIA_PDF_PATH, help="Path to criteria PDF file")
    parser.add_argument("--pdf-directory", default=PDF_DIRECTORY_PATH, help="Directory containing PDFs to compare")
    parser.add_argument("--api-key", help="Claude API key (or set ANTHROPIC_API_KEY env var)")
    parser.add_argument("--csv-output", help="Output CSV file path (default: results.csv)")
    parser.add_argument("--json-output", help="Output JSON file path")
    
    args = parser.parse_args()
    
    # Get API key
    api_key = args.api_key or os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("Please provide API key via --api-key or ANTHROPIC_API_KEY environment variable")
    
    # Set default output paths
    csv_output = args.csv_output or "results.csv"
    
    # Create tool and run comparison
    tool = PDFComparisonTool(api_key)
    
    async def run_comparison():
        results = await tool.compare_pdfs(
            args.criteria_pdf,
            args.pdf_directory,
            csv_output,
            args.json_output
        )
        tool.print_summary(results)
        
        print(f"\n📊 Results saved to:")
        print(f"   CSV: {csv_output}")
        if args.json_output:
            print(f"   JSON: {args.json_output}")
        
        return results
    
    # Run the comparison
    results = asyncio.run(run_comparison())
    
    return results

if __name__ == "__main__":
    main()
